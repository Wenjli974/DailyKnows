#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
新闻简报生成脚本 - 将JSON数据转换为Word文档格式
为CEO准备的精美新闻简报文档
"""

import json
import os
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email.utils import formatdate
from email.encoders import encode_base64
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from dotenv import load_dotenv

def load_news_data(file_path):
    """加载JSON新闻数据文件"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception as e:
        print(f"加载JSON文件时出错：{e}")
        return None

def create_element(name):
    """创建XML元素"""
    return OxmlElement(name)

def create_attribute(element, name, value):
    """设置XML元素属性"""
    element.set(qn(name), value)

def add_page_number(run):
    """添加页码"""
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')
    
    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"
    
    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def set_font(run):
    """设置微软雅黑字体"""
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def create_news_brief(news_data, json_date_str):
    """创建新闻简报Word文档"""
    doc = Document()
    
    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
    
    # 添加页眉
    header = sections[0].header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header_para.add_run('机密文件')
    header_run.font.size = Pt(9)
    header_run.font.color.rgb = RGBColor(0, 0, 0)
    set_font(header_run)
    
    # 添加页脚
    footer = sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(f'第')
    footer_run.font.size = Pt(9)
    set_font(footer_run)
    add_page_number(footer_run)
    footer_run = footer_para.add_run(f'页 | {json_date_str}')
    footer_run.font.size = Pt(9)
    set_font(footer_run)
    
    # 添加标题
    title = doc.add_paragraph()
    title_run = title.add_run(f"{json_date_str} 新闻简报")
    title_run.font.size = Pt(11)
    title_run.font.bold = True
    set_font(title_run)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #title.space_after = Pt(20)
    
    
    # 根据category分类整理新闻
    china_news = [news for news in news_data if news['category'] == "中国新闻"]
    international_news = [news for news in news_data if news['category'] == "国际新闻"]
    auto_news = [news for news in news_data if news['category'] == "汽车相关"]
    
    # 添加第一部分：国内新闻
    add_news_section(doc, "一: 综合热点新闻(国内)", china_news)
    
    # 添加第二部分：国际新闻
    add_news_section(doc, "二: 综合热点新闻(国际)", international_news)
    
    # 添加第三部分：汽车类新闻
    add_news_section(doc, "三: 汽车类热点新闻", auto_news)
    
    return doc

def add_news_section(doc, section_title, news_list):
    """添加新闻部分"""
    # 添加部分标题
    section_heading = doc.add_paragraph()
    section_heading_run = section_heading.add_run(section_title)
    section_heading_run.font.size = Pt(11)
    section_heading_run.font.bold = True
    set_font(section_heading_run)
    section_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #section_heading.space_before = Pt(18)
    #section_heading.space_after = Pt(18)
    
    
    # 添加新闻内容
    for news in news_list:
        # 添加标题
        title_p = doc.add_paragraph()
        title_p.paragraph_format.space_after = Pt(0)
        title_run = title_p.add_run(f"标题：{news['title']}")
        title_run.font.size = Pt(9)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 0, 0)
        set_font(title_run)
        
        # 添加来源
        source_p = doc.add_paragraph()
        source_p.paragraph_format.space_after = Pt(0)
        source_run = source_p.add_run(f"来源：{news['web']}")
        source_run.font.size = Pt(9)
        source_run.font.bold = True
        source_run.font.color.rgb = RGBColor(0, 0, 0)
        set_font(source_run)
        
        # 添加摘要
        summary_p = doc.add_paragraph()
        #summary_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        summary_p.paragraph_format.space_after = Pt(0)
        summary_run = summary_p.add_run(f"摘要：{news['summary']}")
        summary_run.font.size = Pt(9)
        set_font(summary_run)
        
        # 添加分隔线（除了最后一条新闻）
        if news != news_list[-1]:
            separator = doc.add_paragraph()
            separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 设置行距为单倍行距（最小）
            separator.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            separator.paragraph_format.line_spacing = 0.0
            separator.paragraph_format.space_before = Pt(0)
            separator.paragraph_format.space_after = Pt(0)
            
            # 使用正确的方式添加段落边框
            p = separator._p  # 获取段落的XML元素
            pPr = p.get_or_add_pPr()  # 获取段落属性
            
            # 创建边框元素
            pBdr = OxmlElement('w:pBdr')
            
            # 创建底边框元素
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), '000000')
            
            # 添加底边框到边框集合
            pBdr.append(bottom)
            
            # 添加边框集合到段落属性
            pPr.append(pBdr)

def send_news_brief_email(file_path, recipients, subject=None, body=None, sender=None, smtp_server=None, smtp_port=None, smtp_user=None, smtp_password=None):
    """发送新闻简报邮件
    
    参数:
        file_path (str): Word文档路径
        recipients (str或list): 收件人邮箱，可以是字符串或列表
        subject (str): 邮件主题，默认为文件名
        body (str): 邮件正文，默认为简单问候
        sender (str): 发件人邮箱
        smtp_server (str): SMTP服务器地址
        smtp_port (int): SMTP服务器端口
        smtp_user (str): SMTP用户名
        smtp_password (str): SMTP密码或授权码
    """
    # 设置默认值
    if subject is None:
        subject = f"每日新闻简报 - {datetime.now().strftime('%Y年%m月%d日')}"
    
    if body is None:
        body = f"Dear：\n\n请查收附件，谢谢！\n\nBest regards, \n\nWenjin"
    
    # 确保recipients是列表
    if isinstance(recipients, str):
        recipients = [recipients]
    
    # 创建邮件对象
    msg = MIMEMultipart()
    msg['From'] = sender
    msg['To'] = ", ".join(recipients)
    msg['Date'] = formatdate(localtime=True)
    msg['Subject'] = subject
    
    # 添加邮件正文
    msg.attach(MIMEText(body, 'plain', 'utf-8'))
    
    # 添加附件
    if os.path.exists(file_path):
        with open(file_path, 'rb') as f:
            # 使用正确的MIME类型处理Word文档
            attachment = MIMEBase('application', 'vnd.openxmlformats-officedocument.wordprocessingml.document')
            attachment.set_payload(f.read())
            encode_base64(attachment)
            filename = os.path.basename(file_path)
            attachment.add_header('Content-Disposition', f'attachment; filename="{filename}"')
            # 添加正确的内容类型头信息
            attachment.add_header('Content-Type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document', name=filename)
            msg.attach(attachment)
    else:
        print(f"附件文件不存在: {file_path}")
        return False
    
    # 发送邮件
    try:
        server = smtplib.SMTP(smtp_server, smtp_port)
        server.ehlo()
        # 部分SMTP服务器不支持TLS，根据需要取消下面注释
        # 对于smtp.163.com端口25，通常不需要TLS
        # server.starttls()
        server.login(smtp_user, smtp_password)
        server.sendmail(sender, recipients, msg.as_string())
        server.close()
        print(f"邮件已成功发送至: {', '.join(recipients)}")
        return True
    except Exception as e:
        print(f"发送邮件时出错: {e}")
        return False

def main():
    # 加载.env文件中的环境变量
    load_dotenv()
    
    # 获取当前日期
    json_date_str = datetime.now().strftime("%Y%m%d")
    output_date_str = datetime.now().strftime("%Y年%m月%d日")
    
    # 设置文件路径
    json_file = f"D:/pythonProject/DailyKnows/materials/Local_news_{json_date_str}_with_summary.json"
    output_file = f"D:/pythonProject/DailyKnows/DailyReport/{output_date_str} 新闻简报.docx"
    
    # 加载新闻数据
    news_data = load_news_data(json_file)
    if not news_data:
        print("无法加载新闻数据，程序退出")
        return
    
    # 创建新闻简报文档
    doc = create_news_brief(news_data, output_date_str)
    
    # 保存文档
    doc.save(output_file)
    print(f"新闻简报已生成：{output_file}")
    
    # 从环境变量获取邮件配置
    email_config = {
        'sender': os.getenv('EMAIL_SENDER'),
        'smtp_server': os.getenv('EMAIL_SMTP_SERVER'),
        'smtp_port': int(os.getenv('EMAIL_SMTP_PORT')),  # 使用int()转换端口号为整数
        'smtp_user': os.getenv('EMAIL_SMTP_USER'),
        'smtp_password': os.getenv('EMAIL_SMTP_PASSWORD'),
        'recipients': os.getenv('EMAIL_RECIPIENTS', '').split(',')  # 分割逗号分隔的收件人列表
    }
    
    # 检查必要的邮件配置是否存在
    required_config = ['sender', 'smtp_server', 'smtp_user', 'smtp_password', 'recipients']
    missing_config = [key for key in required_config if not email_config[key]]
    
    if missing_config:
        print(f"缺少必要的邮件配置项: {', '.join(missing_config)}")
        print("请在.env文件中补充这些配置项")
        return
    
    #发送邮件
    send_news_brief_email(
        file_path=output_file,
        recipients=email_config['recipients'],
        sender=email_config['sender'],
        smtp_server=email_config['smtp_server'],
        smtp_port=email_config['smtp_port'],
        smtp_user=email_config['smtp_user'],
        smtp_password=email_config['smtp_password']
    )

if __name__ == "__main__":
    main() 