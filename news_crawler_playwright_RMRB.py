from playwright.sync_api import sync_playwright
from datetime import datetime
import docx
import os
import time
import re
import sys
import json

def get_today_date():
    """获取今日日期，格式为YYYYMMDD、YYYYMM和DD"""
    today = datetime.now()
    yyyymmdd = today.strftime("%Y%m%d")
    yyyymm = today.strftime("%Y%m")
    dd = today.strftime("%d")
    return yyyymmdd, yyyymm, dd

def save_to_json(json_path, news_items):
    """将新闻内容保存为JSON格式，如文件已存在则追加内容"""
    json_data = []
    
    # 检查文件是否存在
    if os.path.exists(json_path):
        try:
            # 读取现有文件内容
            with open(json_path, 'r', encoding='utf-8') as f:
                json_data = json.load(f)
                print(f"读取到现有JSON文件: {json_path}，包含 {len(json_data)} 条记录")
                
            # 获取最大ID
            max_id = max([item.get('id', 0) for item in json_data]) if json_data else 0
            
            # 追加新内容
            for i, news in enumerate(news_items, 1):
                json_data.append({
                    "id": max_id + i,
                    "title": news["title"],
                    "source": news["url"],
                    "content": news["content"]
                })
            
            print(f"追加了 {len(news_items)} 条新闻到现有文件")
        except Exception as e:
            print(f"读取现有JSON文件出错: {e}，将创建新文件")
            json_data = []
            for i, news in enumerate(news_items, 1):
                json_data.append({
                    "id": i,
                    "title": news["title"],
                    "source": news["url"],
                    "content": news["content"]
                })
    else:
        # 文件不存在，创建新数据
        for i, news in enumerate(news_items, 1):
            json_data.append({
                "id": i,
                "title": news["title"],
                "source": news["url"],
                "content": news["content"]
            })
    
    # 保存到JSON文件
    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(json_data, f, ensure_ascii=False, indent=4)
    
    print(f"JSON文件已保存：{json_path}，共 {len(json_data)} 条记录")

def append_to_docx(doc_path, news_items):
    """将新闻内容附加到已有的docx文档中"""
    try:
        # 尝试打开现有文档
        doc = docx.Document(doc_path)
    except:
        # 如果文档不存在，创建新文档
        doc = docx.Document()
        doc.add_heading(f'人民日报新闻摘要 - {datetime.now().strftime("%Y年%m月%d日")}', 0)
    
    # 添加分隔线
    doc.add_paragraph('=' * 50)
    doc.add_paragraph(f'更新时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    
    # 添加新闻内容
    for i, news in enumerate(news_items, 1):
        doc.add_heading(f'新闻{i}:', level=1)
        doc.add_paragraph(f'新闻标题：{news["title"]}')
        doc.add_paragraph(f'来源URL：{news["url"]}')
        doc.add_paragraph('新闻主要内容：')
        
        # 添加文章内容段落
        for paragraph in news["content"].split('\n'):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
    
    # 保存文档
    doc.save(doc_path)
    print(f"文档已更新：{doc_path}")

def crawl_people_daily(news_count=3):
    """爬取人民日报头版新闻
    
    参数:
        news_count (int): 要获取的新闻数量，默认为3
    """
    yyyymmdd, yyyymm, dd = get_today_date()
    docx_path = f"materials/Local_news_{yyyymmdd}.docx"
    json_path = f"materials/Local_news_{yyyymmdd}.json"
    rmrb_url = f"http://paper.people.com.cn/rmrb/pc/layout/{yyyymm}/{dd}/node_01.html"
    
    print(f"今日日期: {yyyymmdd}, URL格式: {rmrb_url}")
    print(f"将获取 {news_count} 条新闻")
    
    news_items = []
    
    with sync_playwright() as p:
        # 启动浏览器，设置超时时间
        browser = p.chromium.launch(headless=False)
        context = browser.new_context(
            viewport={'width': 1920, 'height': 1080},
            user_agent='Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        )
        page = context.new_page()
        page.set_default_timeout(30000)  # 将默认超时设置为30秒
        
        try:
            # 打开人民日报网站
            print(f"正在打开人民日报网站: {rmrb_url}")
            page.goto(rmrb_url, wait_until='networkidle')
            time.sleep(3)  # 额外等待以确保页面完全加载
            
            # 截取首页
            homepage_path = f"img/rmrb_homepage_{datetime.now().strftime('%H%M%S')}.png"
            page.screenshot(path=homepage_path, full_page=True)
            print(f"首页截图已保存: {homepage_path}")
            
            # 获取页面HTML源代码，用于调试
            html_source = page.content()
            with open(f"img/rmrb_source_{datetime.now().strftime('%H%M%S')}.html", "w", encoding="utf-8") as f:
                f.write(html_source)
            print("页面源代码已保存")
            
            # 从首页右侧获取新闻列表
            print(f"从首页右侧获取 {news_count} 条新闻列表...")
            
            # 尝试获取右侧新闻列表 - 使用多种可能的选择器
            sidebar_selectors = [
                '.news-list li a', 
                '.right-main a', 
                '.main .right a',
                '.news li a',
                '.sidebar a',
                '.news a',
                '.swiper-container a'
            ]
            
            news_links = []
            
            for selector in sidebar_selectors:
                print(f"尝试获取右侧新闻选择器: {selector}")
                try:
                    # 检查选择器是否有匹配元素
                    count = page.evaluate(f'document.querySelectorAll("{selector}").length')
                    print(f"  - 找到 {count} 个匹配元素")
                    
                    if count > 0:
                        # 获取链接
                        links = page.evaluate(f"""
                        () => {{
                            const elements = document.querySelectorAll('{selector}');
                            return Array.from(elements)
                                .filter(a => a.textContent.trim().length > 3)
                                .slice(0, {news_count})
                                .map(a => ({{
                                    href: a.href,
                                    title: a.textContent.trim() || a.title || "无标题"
                                }}));
                        }}
                        """)
                        
                        if links and len(links) > 0:
                            news_links = links
                            print(f"  - 成功找到新闻链接，使用选择器: {selector}")
                            break
                except Exception as e:
                    print(f"  - 选择器 {selector} 出错: {e}")
            
            # 如果所有选择器都失败，尝试获取所有可见的链接
            if not news_links or len(news_links) == 0:
                print("所有选择器都失败，尝试获取所有可见链接...")
                
                try:
                    # 获取所有可见链接并过滤
                    all_links = page.evaluate(f"""
                    () => {{
                        return Array.from(document.querySelectorAll('a[href]'))
                            .filter(a => {{
                                // 过滤掉不可见的链接
                                const rect = a.getBoundingClientRect();
                                const style = window.getComputedStyle(a);
                                const text = a.textContent.trim();
                                
                                // 确保链接可见、有文本内容、不是导航链接
                                return rect.width > 0 && 
                                       rect.height > 0 && 
                                       style.display !== 'none' && 
                                       style.visibility !== 'hidden' &&
                                       text.length > 5 &&
                                       !a.href.includes('javascript:') &&
                                       !a.href.includes('layout') && // 排除页面导航链接
                                       !a.href.includes('index');
                            }})
                            .map(a => ({{
                                href: a.href,
                                title: a.textContent.trim(),
                                x: a.getBoundingClientRect().x,
                                y: a.getBoundingClientRect().y
                            }}))
                            // 优先考虑页面右侧的链接（x坐标大的）
                            .sort((a, b) => b.x - a.x)
                            .slice(0, {news_count * 2}); // 获取较多链接，后续可进一步筛选
                    }}
                    """)
                    
                    print(f"找到 {len(all_links)} 个可见链接:")
                    for i, link in enumerate(all_links):
                        print(f"{i+1}. {link['title']} (坐标: x={link['x']}, y={link['y']})")
                    
                    # 选择前news_count个链接
                    news_links = []
                    for link in all_links[:news_count]:
                        news_links.append({
                            'href': link['href'],
                            'title': link['title']
                        })
                except Exception as e:
                    print(f"获取可见链接时出错: {e}")
            
            print(f"最终获取到 {len(news_links)} 条新闻链接:")
            for i, link in enumerate(news_links, 1):
                print(f"{i}. 标题: {link['title']}")
                print(f"   URL: {link['href']}")
            
            # 访问每一条新闻并获取内容
            for i, link in enumerate(news_links, 1):
                try:
                    print(f"\n正在处理第 {i} 条新闻: {link['title']}")
                    page.goto(link['href'], wait_until='networkidle')
                    time.sleep(3)  # 额外等待
                    
                    # 截取新闻页面
                    news_screenshot_path = f"img/rmrb_news_{i}_{datetime.now().strftime('%H%M%S')}.png"
                    page.screenshot(path=news_screenshot_path, full_page=True)
                    print(f"新闻页面截图已保存: {news_screenshot_path}")
                    
                    # 保存HTML源码
                    news_html = page.content()
                    with open(f"img/rmrb_news_{i}_source_{datetime.now().strftime('%H%M%S')}.html", "w", encoding="utf-8") as f:
                        f.write(news_html)
                    
                    # 尝试多种选择器获取文章内容
                    content_selectors = [
                        '#ozoom',
                        '.text_c, .content, article',
                        '#artibody',
                        '.article',
                        '.cont_box',
                        '.cnt_bd',
                        '.main-text',
                    ]
                    
                    content = ""
                    for selector in content_selectors:
                        try:
                            if page.locator(selector).count() > 0:
                                print(f"使用内容选择器: {selector}")
                                content = page.evaluate(f"""
                                () => {{
                                    const element = document.querySelector('{selector}');
                                    if (element) {{
                                        return element.innerText;
                                    }}
                                    return "";
                                }}
                                """)
                                
                                if content and len(content) > 100:
                                    print(f"找到内容，长度: {len(content)} 字符")
                                    break
                        except Exception as e:
                            print(f"选择器 {selector} 获取内容出错: {e}")
                    
                    # 如果所有选择器都失败，尝试获取所有段落
                    if not content or len(content) < 100:
                        print("使用备用方法获取内容: 提取所有段落...")
                        try:
                            paragraphs = page.evaluate("""
                            () => {
                                return Array.from(document.querySelectorAll('p'))
                                    .filter(p => {
                                        const text = p.textContent.trim();
                                        const style = window.getComputedStyle(p);
                                        return text.length > 30 && 
                                               style.display !== 'none' && 
                                               style.visibility !== 'hidden';
                                    })
                                    .map(p => p.textContent.trim())
                                    .join('\\n\\n');
                            }
                            """)
                            
                            if paragraphs and len(paragraphs) > 100:
                                content = paragraphs
                                print(f"通过段落获取到内容，长度: {len(content)} 字符")
                        except Exception as e:
                            print(f"获取段落时出错: {e}")
                    
                    # 如果仍然没有内容，使用整个页面文本
                    if not content or len(content) < 100:
                        print("使用备用方法：提取整个页面文本...")
                        try:
                            full_text = page.evaluate("""
                            () => {
                                return document.body.innerText;
                            }
                            """)
                            
                            content = full_text
                            print(f"获取到整个页面文本，长度: {len(content)} 字符")
                        except Exception as e:
                            print(f"获取页面文本时出错: {e}")
                    
                    # 保存新闻信息
                    news_items.append({
                        "title": link['title'],
                        "url": link['href'],
                        "content": content or "无法获取内容"
                    })
                    
                    print(f"第 {i} 条新闻内容获取成功")
                    
                    # 短暂暂停，避免过快请求
                    time.sleep(2)
                    
                except Exception as e:
                    print(f"处理第 {i} 条新闻时出错: {e}")
                    # 添加一个空内容的条目，以保持数量一致
                    news_items.append({
                        "title": link['title'],
                        "url": link['href'],
                        "content": f"获取内容时出错: {str(e)}"
                    })
            
        except Exception as e:
            print(f"爬取过程中出错: {e}")
        finally:
            # 关闭浏览器
            browser.close()
    
    # 将新闻内容保存为JSON格式
    if news_items:
        save_to_json(json_path, news_items)
        print(f"成功获取 {len(news_items)} 条新闻并保存到JSON文件: {json_path}")
    else:
        print("未能获取任何新闻内容")
    
    return len(news_items) > 0

if __name__ == "__main__":
    # 确保img目录存在
    os.makedirs("materials", exist_ok=True)
    
    # 设置要获取的新闻数量
    # 默认获取3条新闻，也可以通过命令行参数修改
    news_count = 1
    
    # 解析命令行参数
    if len(sys.argv) > 1:
        try:
            news_count = int(sys.argv[1])
            print(f"通过命令行参数设置获取 {news_count} 条新闻")
        except ValueError:
            print(f"无效的命令行参数: {sys.argv[1]}，使用默认值3")
    
    try:
        success = crawl_people_daily(news_count)
        if success:
            print("程序执行成功!")
        else:
            print("程序执行完成，但未能获取新闻内容")
    except Exception as e:
        print(f"程序执行出错: {e}")